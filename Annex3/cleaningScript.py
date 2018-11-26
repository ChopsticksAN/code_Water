import docx
import xlwt
from docx import Document

def getText(filename):
    doc = docx.Document(filename)
    fullText=[]
    for i in doc.paragraphs:
        fullText.append(i.text)
    return fullText


def cleanFulltext(fullText):
    for i in range(len(fullText), -1, -1):
        if i >= len(fullText):
            continue
        elif len(fullText[i]) < 6:
            fullText.remove(fullText[i])
        elif fullText[i][5] != '2':
            fullText.remove((fullText[i]))

    for i in range(0, len(fullText)):
        fullText[i] = fullText[i][5:]
    return fullText

def writeEndTable(end_table):

    return

def main():
    path = ".\\origin.doc"
    document = Document(path)
    fullText = getText(path)
    fullText = cleanFulltext(fullText)

    #print(len(fullText))

    #for i in range(0, len(fullText)):
    #    print(fullText[i][5:])

    workbook = xlwt.Workbook(encoding='ascii')
    worksheet = workbook.add_sheet('CleanData')

    tables = document.tables
    t = len(tables)
    print(t)

    rows = 0
    cols = 0
    ele = 0

    #print(fullText[0])
    #print(len(fullText[0]))
    for j in range(3, 9):
        s = tables[0].cell(1, j).text
        worksheet.write(rows, cols, label=s)
        cols += 1

    worksheet.write(rows, cols, label="水流量")
    worksheet.write(rows, cols + 1, label="水流速")

    cols = 0
    rows += 2
    for p in range(0, t - 1):
        worksheet.write(rows, cols, label=fullText[ele])
        rows += 1
        ele += 1
        table = tables[p]
        for i in range(2, len(table.rows)):
            cols = 0
            for j in range(3, 9):
                s = table.cell(i, j).text
                worksheet.write(rows, cols, label=s)
                cols += 1
            worksheet.write(rows, cols, label="0")
            worksheet.write(rows, cols + 1, label="0")
            rows += 1
        rows += 1
        cols = 0

    end_table = tables[t - 1]
    writeEndTable(end_table)

   # for i in range(0, 10):
       # print(worksheet.row(1)[0].value)

    workbook.save(".\\CleanedData.xls")

if __name__=="__main__":
    main()




