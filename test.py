import docx
import xlwt
from docx import Document

path = "C:\\Users\\王金宇\\Desktop\\数学建模\\111.doc"
document = Document(path)

'''
for paragraph in document.paragraphs:
    print(paragraph.text)
'''

workbook = xlwt.Workbook(encoding='ascii')
worksheet = workbook.add_sheet('My_Worksheet')

tables = document.tables
t = len(tables)
print(t)

rows = 0
cols = 0

for p in range(0, t):
    table = tables[p]
    for i in range(1, len(table.rows)):
        #result = ""
        cols = 0
        for j in range(3, 9):
            s = table.cell(i, j).text
            #print(s);
            worksheet.write(rows, cols, label=s)
            cols += 1
        rows += 1
    rows += 1

workbook.save('C:\\Users\\王金宇\\Desktop\\数学建模\\111.xls')

'''
            if i == 1:
                result = result + " " + s
            else:
                result += s;
            print(result)
            '''


